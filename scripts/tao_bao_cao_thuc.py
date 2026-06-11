"""
Báo cáo THỰC TẾ — Ngân hàng & BĐS Phía Nam
Dữ liệu: RSS + WebSearch, ngày 05/06/2026
Chạy: python scripts/tao_bao_cao_thuc.py
Output: outputs/BaoCao_NganHang_BDS_PhiaNam_THUCTE_20260605.docx
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
from pathlib import Path


# ── Helpers ────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_heading(doc, text, level=1, color='1A3A5C'):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = RGBColor.from_string(color)
    return p


def add_para(doc, text, bold=False, size=11, color=None, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return p


def add_bullet(doc, text, size=11, bold=False):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    return p


def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'D5E8F5')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def build_signal_table(doc, data):
    table = doc.add_table(rows=1, cols=len(data))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    row = table.rows[0]
    for i, (label, value, color) in enumerate(data):
        cell = row.cells[i]
        set_cell_bg(cell, color)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run(value + '\n')
        r1.bold = True
        r1.font.size = Pt(13)
        r1.font.color.rgb = RGBColor(255, 255, 255)
        r2 = p.add_run(label)
        r2.font.size = Pt(9)
        r2.font.color.rgb = RGBColor(220, 235, 250)
    for cell in row.cells:
        tc = cell._tc
        trPr = tc.getparent().get_or_add_trPr()
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), '700')
        trPr.append(trHeight)
    return table


def build_data_table(doc, headers, rows_data, col_widths=None):
    table = doc.add_table(rows=1 + len(rows_data), cols=len(headers))
    table.style = 'Table Grid'
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, '1A3A5C')
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(255, 255, 255)
    for r_idx, row_data in enumerate(rows_data):
        row = table.rows[r_idx + 1]
        bg = 'EBF5FB' if r_idx % 2 == 0 else 'FFFFFF'
        for c_idx, (cell_val, is_center, is_bold) in enumerate(row_data):
            cell = row.cells[c_idx]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if is_center else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(cell_val))
            run.font.size = Pt(10)
            run.bold = is_bold
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table


# ── Build report ──────────────────────────────────────────────────────────

def build_report():
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.0)
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal'].font.size = Pt(11)

    # ── TIÊU ĐỀ ────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run('BÁO CÁO THỊ TRƯỜNG')
    r.bold = True; r.font.size = Pt(20)
    r.font.color.rgb = RGBColor.from_string('1A3A5C')

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(4)
    r2 = p2.add_run('NGÂN HÀNG & BẤT ĐỘNG SẢN PHÍA NAM')
    r2.bold = True; r2.font.size = Pt(16)
    r2.font.color.rgb = RGBColor.from_string('2471A3')

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_after = Pt(14)
    r3 = p3.add_run('Ngày 05/06/2026  |  TP.HCM · Bình Dương · Đồng Nai · Long An · Bà Rịa-Vũng Tàu')
    r3.font.size = Pt(10)
    r3.font.color.rgb = RGBColor.from_string('7F8C8D')
    r3.italic = True
    add_divider(doc)

    # ── BẢNG TÍN HIỆU ──────────────────────────────────────────────
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    build_signal_table(doc, [
        ('Xu hướng Lãi suất',   '↑ TĂNG NHẸ',   '2471A3'),
        ('Tín dụng BĐS',        '→ KIỂM SOÁT',   '117864'),
        ('Tín hiệu Thị trường', '🟡 THẬN TRỌNG', 'B7770D'),
        ('Bài phân tích hôm nay', '10 tin', '6C3483'),
    ])
    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # ── I. TỔNG QUAN ───────────────────────────────────────────────
    add_heading(doc, 'I. TỔNG QUAN THỊ TRƯỜNG — 05/06/2026', level=1)
    add_para(doc, 'Điểm nổi bật hôm nay:', bold=True)
    add_bullet(doc, 'Lãi suất kỳ hạn ngắn (1-3 tháng) tăng lên mức trần 4,75%/năm tại TPBank; nhiều ngân hàng thực trả trên 8%/năm dù niêm yết 7%.')
    add_bullet(doc, 'Giá căn hộ TP.HCM chạm mốc 190 triệu đồng/m² — lần đầu vượt nhà liền thổ. Trong khi đó nhà riêng/nhà phố đang giảm giá do khó thanh khoản.')
    add_bullet(doc, 'BĐS phía Nam "nóng bên cung, trầm lắng bên cầu": nguồn cung phục hồi mạnh nhưng sức mua yếu, giao dịch căn hộ vẫn chờ khớp lệnh.')
    add_bullet(doc, 'NHNN kiểm soát tín dụng BĐS — tốc độ tăng tín dụng BĐS của mỗi TCTD không được vượt tốc độ tăng chung (mục tiêu 15% toàn hệ thống).')
    add_bullet(doc, 'Lãi suất liên ngân hàng VND quay đầu giảm sau khi tăng mạnh gần 11%; USD tăng giá kịch trần do căng thẳng Trung Đông leo thang.')
    doc.add_paragraph()

    # ── II. LÃI SUẤT ───────────────────────────────────────────────
    add_heading(doc, 'II. XU HƯỚNG LÃI SUẤT', level=1)
    add_para(doc, 'A. Lãi suất huy động tháng 6/2026', bold=True, color='1A3A5C')

    rows_hd = [
        [('Vietcombank', False, False), ('4,75%', True, False), ('5,9%', True, False), ('5,9%', True, False), ('→ Ổn định', True, False)],
        [('BIDV',        False, False), ('4,75%', True, False), ('5,9%', True, False), ('5,9%', True, False), ('→ Ổn định', True, False)],
        [('VietinBank',  False, False), ('4,75%', True, False), ('5,9%', True, False), ('5,9%', True, False), ('→ Ổn định', True, False)],
        [('Techcombank', False, False), ('~5,0%', True, False), ('~6,0%', True, False), ('~6,5%', True, False), ('→ Ổn định', True, False)],
        [('MB Bank',     False, False), ('5,0%',  True, False), ('6,5%',  True, False), ('7,0%', True, True),  ('↑ Tăng',   True, True)],
        [('ACB',         False, False), ('~4,9%', True, False), ('~6,2%', True, False), ('~6,5%', True, False), ('→ Ổn định', True, False)],
        [('TPBank',      False, False), ('4,75%', True, True),  ('~6,5%', True, False), ('~7,0%', True, False), ('↑ Tăng kỳ ngắn', True, True)],
        [('Cake/VPBank (online)', False, False), ('—', True, False), ('7,4%', True, True), ('7,4%', True, True), ('↑ Cao nhất online', True, True)],
        [('HLBank',      False, False), ('—', True, False),     ('7,3%', True, True),  ('—', True, False),      ('Cao nhất tại quầy', True, True)],
    ]
    build_data_table(doc,
        ['Ngân hàng', 'Kỳ hạn 1-3T', 'Kỳ hạn 12T', 'Kỳ hạn 24T+', 'Xu hướng'],
        rows_hd, col_widths=[4.5, 2.8, 2.8, 2.8, 3.5])

    doc.add_paragraph()
    add_para(doc, 'Ghi chú: Lãi suất liên ngân hàng VND quay đầu giảm sau khi tăng mạnh gần 11%/năm (04/06). '
             'Nhiều ngân hàng thực trả trên 8%/năm dù niêm yết 7% — cho thấy áp lực cạnh tranh huy động vốn.',
             size=10, color='7F8C8D', space_after=10)

    add_para(doc, 'B. Lãi suất cho vay mua BĐS (mức hiện hành)', bold=True, color='1A3A5C')

    rows_vay = [
        [('Vietcombank', False, False), ('9,6%', True, False), ('9,9%', True, False), ('13,6%', True, False),  ('70-80%', True, False), ('25 năm', True, False)],
        [('BIDV',        False, False), ('9,7%', True, False), ('10,1%', True, False), ('13,5%', True, False), ('70%',    True, False), ('25 năm', True, False)],
        [('VietinBank',  False, False), ('10,0%', True, False), ('10,0%', True, False), ('—', True, False),    ('70%',    True, False), ('25 năm', True, False)],
        [('Techcombank', False, False), ('8,5%', True, True),  ('9,5%', True, False),  ('—', True, False),     ('70-80%', True, False), ('25 năm', True, False)],
        [('VIB',         False, False), ('9,9%', True, False), ('12,0%', True, False), ('—', True, False),     ('70%',    True, False), ('20 năm', True, False)],
        [('ACB',         False, False), ('9,5%', True, False), ('10,5%', True, False), ('—', True, False),     ('70%',    True, False), ('25 năm', True, False)],
        [('NHTM Nhà nước (NƠXH)', False, True), ('4,6%', True, True), ('—', True, False), ('—', True, False), ('70%',    True, False), ('15 năm', True, False)],
        [('Agribank (NTN <35T)', False, False), ('5,6%', True, True), ('—', True, False), ('—', True, False),  ('70%',    True, False), ('15 năm', True, False)],
    ]
    build_data_table(doc,
        ['Ngân hàng', 'Ưu đãi 6T đầu', 'Ưu đãi 12T đầu', 'Sau ưu đãi', 'LTV tối đa', 'Thời hạn'],
        rows_vay, col_widths=[4.5, 2.5, 2.5, 2.5, 2.5, 2.2])

    doc.add_paragraph()
    add_para(doc, 'Nhận định: Mặt bằng lãi suất cho vay mua nhà đang ở mức cao nhất trong 2 năm qua. '
             'Người mua nhà cần tính toán kỹ khả năng trả nợ, đặc biệt sau giai đoạn ưu đãi.',
             size=10, color='922B21', bold=True, space_after=10)

    # ── III. TÍN DỤNG BĐS ──────────────────────────────────────────
    add_heading(doc, 'III. TÍN DỤNG BẤT ĐỘNG SẢN', level=1)
    add_para(doc, 'A. Chính sách NHNN 2026', bold=True, color='1A3A5C')
    add_bullet(doc, 'Mục tiêu tăng trưởng tín dụng toàn hệ thống: 15% năm 2026.')
    add_bullet(doc, 'Quy định mới: Tốc độ tăng tín dụng BĐS của từng TCTD không được vượt tốc độ tăng tín dụng chung của chính TCTD đó. Vi phạm → NHNN giảm room tăng trưởng tín dụng.')
    add_bullet(doc, 'NHNN không siết cho vay BĐS nhưng kiểm soát chặt — tức giữ nguyên hệ số rủi ro, không nới thêm.')
    add_bullet(doc, 'Thông tư 07/2026/TT-NHNN (hiệu lực 20/6/2026): Sửa đổi quy định về môi giới tiền tệ — yêu cầu kiểm soát rủi ro giao dịch tài chính.')

    doc.add_paragraph()
    add_para(doc, 'B. Tín hiệu đáng chú ý từ hệ thống ngân hàng', bold=True, color='1A3A5C')
    add_bullet(doc, 'Agribank: Triển khai gói tín dụng xanh 3.000 tỷ đồng cho khách hàng cá nhân — cơ hội vay ưu đãi gắn BĐS xanh/nhà ở NƠXH.')
    add_bullet(doc, 'Sacombank: Siết nợ, thu hồi hơn 500 sổ đỏ tại Viva City (Đồng Nai/LDG) — cảnh báo rủi ro pháp lý với các dự án có vướng nợ ngân hàng.')
    add_bullet(doc, 'Một ngân hàng đấu giá khoản nợ hơn 730 tỷ của chủ đầu tư bệnh viện lớn — nợ xấu BĐS phi nhà ở đang được xử lý.')
    add_bullet(doc, '7 ngân hàng sở hữu tài sản trên 1 triệu tỷ đồng — năng lực cho vay BĐS tập trung tại nhóm lớn.')
    add_bullet(doc, 'Quy định mới: Tài khoản ngủ đông 3 năm không giao dịch có thể bị đóng — ảnh hưởng tài khoản phong tỏa trong giao dịch BĐS.')

    doc.add_paragraph()

    # ── IV. THỊ TRƯỜNG PHÍA NAM ─────────────────────────────────────
    add_heading(doc, 'IV. PHÂN TÍCH THỊ TRƯỜNG PHÍA NAM', level=1)
    add_para(doc, 'Tổng quan: "Nóng bên cung, trầm lắng bên cầu" (VnExpress, 01/06/2026)', bold=True, color='922B21')
    add_para(doc, 'Nguồn cung phục hồi mạnh sau 2 năm tắc nghẽn pháp lý, nhưng sức mua thực tế còn yếu do lãi suất vay cao và thu nhập người mua chưa bắt kịp giá. '
             'Thị trường đang tự sàng lọc: dự án có pháp lý sạch, vị trí tốt vẫn giao dịch được; dự án vướng mắc hoặc giá cao so với nhu cầu thực đang tồn kho.', space_after=10)

    khu_vuc_data = [
        ('TP.HCM', '1A3A5C', [
            'Giá căn hộ: Chạm mốc 190 triệu đồng/m² — lần đầu vượt nhà liền thổ (VnExpress, 03/06/2026).',
            'Nhà riêng, nhà phố: Đang giảm giá do khó bán — thanh khoản thấp nhất 5 quý gần đây.',
            'Giao dịch căn hộ: Sau thời gian tăng nóng, đà tăng giá chững lại, giao dịch "vẫn chờ khớp lệnh".',
            'Chung cư cũ bám dự án hạng sang để neo giá — tạo mặt bằng giá ảo tại nhiều khu vực.',
            'TP.HCM đề xuất hệ số K điều chỉnh giá đất cao nhất 2,59 lần → tiền sử dụng đất dự kiến tăng cao — áp lực lên giá bán mới.',
            'Hai dự án hơn 600 ha (Bình Quới-Thanh Đa) được yêu cầu bàn giao mặt bằng tháng 10/2026 — nguồn cung lớn sắp vào thị trường.',
            'Dự án nổi bật mở bán: The Esme hơn 4,3 ha phía Đông TP.HCM (Quốc lộ 1K, 05/06/2026).',
            'Chỉ số sản xuất công nghiệp (IIP) TP.HCM tăng cao nhất 5 năm → hỗ trợ nhu cầu nhà ở thực.',
            'TP.HCM tầm nhìn 100 năm: Định hướng trở thành top 100 đô thị chất lượng sống tốt nhất thế giới năm 2050.',
        ]),
        ('Bình Dương', '117864', [
            'Phân khúc cao cấp chiếm 60% nguồn cung mới — mức cao chưa từng có trong lịch sử Bình Dương.',
            'Giá căn hộ sơ cấp tăng 14-15%/năm — vượt tốc độ tăng thu nhập người mua.',
            'Becamex (chủ đầu tư lớn nhất) đề xuất giảm sở hữu nhà nước — cơ cấu lại mô hình phát triển đô thị.',
            'TP.HCM muốn kéo tuyến metro từ Thủ Dầu Một thẳng ga Tao Đàn → nếu thực hiện, kết nối BD-TPHCM sẽ đột phá, hỗ trợ BĐS vùng phụ cận.',
            'Nguy cơ giải phóng mặt bằng chậm gây trễ 30-40% dự án mới.',
        ]),
        ('Đồng Nai', '6E2F8D', [
            'Izumi City (Nam Long - Đồng Nai): Thu hút nhà đầu tư phía Bắc quan tâm mô hình đô thị tích hợp.',
            'Vinhomes Golden City: Được quan tâm nhờ hạ tầng giao thông, khu công nghiệp phát triển quanh dự án.',
            'Kim Oanh Group (quỹ đất lớn tại Đồng Nai): Đang triển khai đa dạng nguồn cung, nhận Giải thưởng Quốc gia BĐS.',
            'Sacombank thu hồi hơn 500 sổ đỏ tại Viva City (LDG) → cảnh báo pháp lý với phân khúc giá rẻ/pháp lý chưa hoàn thiện.',
            'Giá sơ cấp căn hộ Đồng Nai tăng 14-15%/năm — tương đồng Bình Dương.',
        ]),
        ('Long An', '935116', [
            'Kim Oanh Group (hơn 1.800 ha quỹ đất, phần lớn tại Long An) đang đẩy mạnh phát triển nhiều dự án.',
            'Thị trường đất nền Long An: Giao dịch ổn định hơn TP.HCM nhờ giá còn thấp hơn mặt bằng chung.',
            'Hạ tầng kết nối TP.HCM-Long An đang cải thiện dần — yếu tố dài hạn quan trọng.',
            'Phân khúc nhà ở giá vừa túi tiền (1-3 tỷ) vẫn là nhu cầu thực lớn chưa được đáp ứng.',
        ]),
        ('Bà Rịa - Vũng Tàu', '922B21', [
            'Nhiệt điện Bà Rịa đảm bảo vận hành, sẵn sàng cung ứng điện mùa cao điểm — ổn định hạ tầng công nghiệp, hỗ trợ BĐS khu công nghiệp.',
            'Thị trường BĐS nghỉ dưỡng tiếp tục trầm lắng theo xu hướng chung toàn quốc.',
            'BĐS khu công nghiệp: Vẫn là điểm sáng dài hạn nhờ thu hút FDI và phát triển cảng biển.',
            'Liên doanh Việt Nam-Singapore-Nhật Bản vừa thành lập phát triển BĐS: Khu vực BRVT nằm trong tầm ngắm.',
        ]),
    ]

    for kv, color, bullets in khu_vuc_data:
        add_para(doc, f'  {kv}', bold=True, size=12, color=color)
        for b in bullets:
            add_bullet(doc, b)
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── V. TIN TỨC NỔI BẬT ─────────────────────────────────────────
    add_heading(doc, 'V. TIN TỨC NỔI BẬT — 05/06/2026', level=1)

    rows_tin = [
        [('Giá căn hộ TP.HCM vượt nhà liền thổ, chạm mốc 190 triệu/m²', False, True),
         ('VnExpress BĐS', True, False), ('→ Ổn định', True, False), ('→ Không đổi', True, False),
         ('Áp lực tâm lý lên NĐT: chung cư cao cấp neo giá mới, nhà phố mất lợi thế tương đối', False, False),
         ('CHỜ ĐỢI', True, True)],

        [('Nhà riêng, nhà phố TP.HCM giảm giá vì khó bán', False, True),
         ('VnExpress', True, False), ('→ Ổn định', True, False), ('→ Không đổi', True, False),
         ('Thanh khoản nhà phố/nhà riêng yếu — cơ hội mặc cả cho người mua ở thực', False, False),
         ('GIỮ/CHỜ', True, False)],

        [('Sacombank thu hồi hơn 500 sổ đỏ tại Viva City của LDG', False, True),
         ('VnExpress', True, False), ('↑ Tăng', True, False), ('→ Thắt chặt', True, True),
         ('Cảnh báo rủi ro pháp lý tại các dự án vướng nợ ngân hàng — kiểm tra kỹ trước khi mua', False, True),
         ('CHỜ ĐỢI', True, True)],

        [('Lãi suất liên ngân hàng VND quay đầu giảm sau khi tăng gần 11%', False, False),
         ('Thanh Niên', True, False), ('↓ Giảm', True, True), ('→ Không đổi', True, False),
         ('Thanh khoản hệ thống cải thiện — có thể hỗ trợ giảm lãi vay trong ngắn hạn', False, False),
         ('THEO DÕI', True, False)],

        [('NHNN kiểm soát tín dụng BĐS: không vượt tốc độ tăng tín dụng chung', False, False),
         ('CafeF', True, False), ('→ Ổn định', True, False), ('→ Kiểm soát', True, True),
         ('Không có room mới cho BĐS — vốn tín dụng không đổ thêm vào BĐS đầu cơ', False, False),
         ('GIỮ/CHỜ', True, False)],

        [('Lãi suất cho vay cao nhất 2 năm: Vietcombank 9,6%; BIDV 9,7% (ưu đãi 6T)', False, True),
         ('VietnamNet', True, False), ('↑ Tăng', True, True), ('→ Không đổi', True, False),
         ('Chi phí vốn cao — ảnh hưởng trực tiếp đến khả năng vay mua nhà của người có thu nhập trung bình', False, True),
         ('CHỜ ĐỢI', True, True)],

        [('BĐS phía Nam "nóng bên cung, trầm lắng bên cầu"', False, True),
         ('VnExpress', True, False), ('→ Ổn định', True, False), ('→ Không đổi', True, False),
         ('Thị trường phân hóa mạnh — chọn lọc dự án pháp lý sạch, giá hợp lý, vị trí tốt', False, False),
         ('GIỮ/CHỜ', True, False)],

        [('Agribank triển khai gói tín dụng xanh 3.000 tỷ đồng', False, False),
         ('VnExpress', True, False), ('↓ Giảm ưu đãi', True, True), ('✅ Nới lỏng', True, True),
         ('Cơ hội vay lãi suất thấp gắn với BĐS xanh, nhà ở xã hội — nhà đầu tư nhu cầu thực cần chú ý', False, True),
         ('MUA VỀ', True, True)],

        [('TP.HCM muốn kéo metro từ Thủ Dầu Một thẳng ga Tao Đàn', False, False),
         ('Thanh Niên', True, False), ('→ Ổn định', True, False), ('→ Không đổi', True, False),
         ('Tích cực dài hạn cho BĐS dọc tuyến BD-TPHCM — chưa tác động ngay trong năm 2026', False, False),
         ('GIỮ/CHỜ', True, False)],

        [('Sản xuất TP.HCM tăng cao nhất 5 năm — IIP tăng mạnh', False, False),
         ('VnExpress', True, False), ('→ Ổn định', True, False), ('→ Không đổi', True, False),
         ('Tín hiệu tích cực: kinh tế thực phục hồi → nhu cầu nhà ở dài hạn được hỗ trợ', False, True),
         ('MUA VỀ', True, True)],
    ]

    table = doc.add_table(rows=1 + len(rows_tin), cols=6)
    table.style = 'Table Grid'

    headers_tin = ['Tiêu đề tin tức', 'Nguồn', 'Lãi suất', 'Tín dụng BĐS', 'Tác động NĐT phía Nam', 'Khuyến nghị']
    hdr_row = table.rows[0]
    for i, h in enumerate(headers_tin):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, '1A3A5C')
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True; run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255, 255, 255)

    kn_colors = {'MUA VỀ': '1E8449', 'GIỮ/CHỜ': 'B7770D', 'CHỜ ĐỢI': '922B21', 'THEO DÕI': '1A3A5C'}
    for r_idx, row_data in enumerate(rows_tin):
        row = table.rows[r_idx + 1]
        bg = 'EBF5FB' if r_idx % 2 == 0 else 'FFFFFF'
        for c_idx, (cell_val, is_center, is_bold) in enumerate(row_data):
            cell = row.cells[c_idx]
            if c_idx == 5:
                kn_color = kn_colors.get(cell_val, '1A3A5C')
                set_cell_bg(cell, kn_color)
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(str(cell_val))
                run.bold = True; run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(255, 255, 255)
            else:
                set_cell_bg(cell, bg)
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if is_center else WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(str(cell_val))
                run.font.size = Pt(9)
                run.bold = is_bold

    col_widths_tin = [4.5, 1.8, 2.0, 2.2, 4.5, 1.8]
    for i, w in enumerate(col_widths_tin):
        for row in table.rows:
            row.cells[i].width = Cm(w)

    doc.add_paragraph()

    # ── VI. PHÂN TÍCH TỔNG QUAN ────────────────────────────────────
    add_heading(doc, 'VI. PHÂN TÍCH TỔNG QUAN', level=1)

    analysis_sections = [
        ('1. Xu hướng Lãi suất',
         'Lãi suất huy động đang leo thang cục bộ — đặc biệt nhóm kỳ hạn ngắn (4,75%/năm, chạm trần) và '
         'các ngân hàng nhỏ/trung (thực trả trên 8%/năm). Điều này phản ánh áp lực thanh khoản cục bộ tại '
         'một số ngân hàng sau giai đoạn giải ngân mạnh. Lãi suất liên ngân hàng VND đã quay đầu giảm sau '
         'khi chạm đỉnh ~11% — tín hiệu NHNN bơm tiền hỗ trợ. Tuy nhiên, lãi suất cho vay mua BĐS tại các '
         'NHTM lớn (9,6-10,5% giai đoạn ưu đãi, 13-14% sau ưu đãi) vẫn ở mức cao nhất 2 năm — chưa có '
         'dấu hiệu giảm trong ngắn hạn. Kỳ vọng lãi suất cho vay giảm nhẹ trong H2/2026 nếu NHNN tiếp tục '
         'hỗ trợ thanh khoản và lạm phát được kiểm soát.'),
        ('2. Tín dụng BĐS',
         'NHNN đang duy trì chính sách "không siết nhưng kiểm soát chặt": tín dụng BĐS không có room '
         'riêng được nới thêm — tốc độ tăng tín dụng BĐS của từng ngân hàng không được vượt tốc độ tăng '
         'tín dụng chung (mục tiêu 15%). Trong thực tế, điều này có nghĩa là dòng vốn tín dụng vào BĐS '
         'sẽ tăng tương ứng với tốc độ tăng tổng tín dụng — không bị cắt đột ngột nhưng cũng không có '
         'cú hích mới. Điểm tích cực: Agribank tung gói tín dụng xanh 3.000 tỷ — mở hướng vay ưu đãi '
         'cho BĐS xanh/NƠXH. Điểm cần thận trọng: Sacombank thu hồi sổ đỏ Viva City, ngân hàng đấu giá '
         'nợ xấu BĐS — phân khúc dự án có vấn đề pháp lý/tài chính đang bị xử lý mạnh.'),
        ('3. Thị trường Phía Nam',
         'TP.HCM: Hai xu hướng trái chiều đang diễn ra đồng thời — căn hộ cao cấp neo giá ở mức mới cao '
         '(190 triệu/m²) trong khi nhà riêng/nhà phố đang chịu áp lực giảm giá do thanh khoản yếu. '
         'Đây là dấu hiệu phân hóa rõ nét: dòng tiền chảy vào căn hộ cao cấp tại vị trí tốt, trong khi '
         'bỏ qua nhà liền thổ ở xa trung tâm. Bình Dương & Đồng Nai: Giá sơ cấp tăng 14-15%/năm nhưng '
         'thanh khoản thứ cấp yếu — nguy cơ bong bóng cục bộ tại phân khúc cao cấp mới. Long An: Điểm '
         'cân bằng tốt hơn — giá đất nền còn dư địa, phù hợp nhu cầu thực dài hạn. BRVT: BĐS công '
         'nghiệp tiếp tục là điểm sáng dài hạn nhờ FDI và hạ tầng cảng biển.'),
        ('4. Rủi ro cần theo dõi',
         '(1) Lãi suất cao kéo dài làm giảm sức mua thực tế — nguy cơ tồn kho tăng tại phân khúc trung-cao cấp. '
         '(2) USD tăng giá kịch trần (căng thẳng Trung Đông) → tỷ giá biến động → NHNN có thể phải hút tiền '
         'đồng để bảo vệ tỷ giá → thanh khoản VND căng thẳng hơn. '
         '(3) Pháp lý dự án: Sacombank thu hồi sổ đỏ Viva City là cảnh báo đỏ — dự án BĐS chưa hoàn thiện '
         'pháp lý hoặc chủ đầu tư có nợ ngân hàng đang bị siết mạnh. '
         '(4) Tiền sử dụng đất tăng theo đề xuất hệ số K mới tại TP.HCM — áp lực chi phí đất sẽ đẩy giá '
         'bán sơ cấp lên trong H2/2026.'),
    ]

    for title, content in analysis_sections:
        add_para(doc, title, bold=True, size=11, color='1A3A5C')
        add_para(doc, content, size=11, space_after=8)

    add_divider(doc)

    # ── VII. KHUYẾN NGHỊ ───────────────────────────────────────────
    add_heading(doc, 'VII. KHUYẾN NGHỊ — NĐT BĐS PHÍA NAM (05/06/2026)', level=1)

    khuyen_nghi_data = [
        ('🟢  NÊN LÀM — Cơ hội', '27AE60', [
            'Tận dụng gói tín dụng xanh 3.000 tỷ của Agribank: Nếu đang có nhu cầu vay mua NƠXH hoặc BĐS xanh, đây là cơ hội lãi suất ưu đãi thấp hơn thị trường đáng kể.',
            'Nhà phố/nhà riêng TP.HCM đang điều chỉnh giá: Người mua ở thực có thể thương lượng giá tốt hơn 5-10% so với đỉnh — nhất là các dự án chủ đầu tư cần dòng tiền.',
            'Theo dõi đất nền Long An giá vừa phải: Phân khúc 500 triệu - 1,5 tỷ, pháp lý rõ ràng, gần trục giao thông lớn — phù hợp đầu tư dài hạn 3-5 năm.',
            'BĐS công nghiệp BRVT và Đồng Nai: Nhu cầu thuê nhà xưởng/kho bãi từ FDI vẫn tăng — phân khúc này ít bị ảnh hưởng bởi lãi suất và mang lại dòng tiền thuê ổn định.',
        ]),
        ('🟡  THEO DÕI — Trung lập', 'B7770D', [
            'Căn hộ cao cấp TP.HCM (150-190 triệu/m²): Giá đã ở vùng đỉnh, thanh khoản thứ cấp yếu — không phù hợp lướt sóng. Phù hợp giữ dài hạn nếu đã mua.',
            'Theo dõi động thái NHNN về tỷ giá và lãi suất điều hành trong tháng 6: Tỷ giá USD tăng mạnh có thể buộc NHNN thắt chặt tiền tệ → ảnh hưởng lãi vay.',
            'Dự án Bình Dương & Đồng Nai phân khúc cao cấp: Giá tăng 14-15%/năm nhưng thanh khoản chưa tương xứng — cần đánh giá kỹ dòng tiền và khả năng thoát hàng.',
            'Metro kéo từ Thủ Dầu Một - Tao Đàn: Dự án chưa phê duyệt chính thức — không nên vào tiền sớm dựa trên thông tin này.',
        ]),
        ('🔴  THẬN TRỌNG — Rủi ro cao', '922B21', [
            'TRÁNH các dự án có chủ đầu tư đang vướng nợ ngân hàng: Vụ Sacombank thu hồi sổ đỏ Viva City là cảnh báo rõ ràng — kiểm tra kỹ pháp lý và tình trạng thế chấp dự án tại ngân hàng trước khi đặt cọc.',
            'KHÔNG dùng đòn bẩy cao (vay >70% giá trị) trong thời điểm lãi suất 9,6-10% + sau ưu đãi lên 13-14%: Gánh nặng lãi suất có thể xói mòn hoàn toàn lợi nhuận đầu tư trong 2-3 năm đầu.',
            'Tránh nhà phố/nhà riêng TP.HCM vùng ven xa trung tâm (>20km) với giá cao: Thanh khoản yếu nhất thị trường hiện tại — nguy cơ chôn vốn kéo dài.',
            'Thận trọng với dự án hứa hẹn hệ số K cao / tiền sử dụng đất thấp bất thường: TP.HCM đang xem xét tăng hệ số K — giá trị thực có thể cao hơn.',
        ]),
    ]

    for title, color_hex, items in khuyen_nghi_data:
        add_para(doc, title, bold=True, size=12, color=color_hex, space_after=4)
        for item in items:
            add_bullet(doc, item)
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── NGUỒN & FOOTER ──────────────────────────────────────────────
    add_divider(doc)
    add_para(doc, 'Nguồn dữ liệu:', bold=True, size=10, color='1A3A5C')
    sources = [
        'VnExpress BĐS — vnexpress.net/bat-dong-san (05/06/2026)',
        'VnExpress Kinh tế — vnexpress.net/kinh-doanh',
        'Thanh Niên Kinh tế — thanhnien.vn/kinh-te',
        'VietnamNet Kinh tế — vietnamnet.vn/kinh-doanh',
        'CafeF — cafef.vn (lãi suất, tín dụng BĐS)',
        'Savills Vietnam — Báo cáo Q1/2026 thị trường TP.HCM',
        'NHNN — Thông tư 07/2026/TT-NHNN (hiệu lực 20/6/2026)',
    ]
    for s in sources:
        add_bullet(doc, s, size=9)

    p_footer = doc.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_footer.paragraph_format.space_before = Pt(12)
    r_footer = p_footer.add_run(
        'Báo cáo nội bộ · Ngân Hàng & BĐS Phía Nam · 05/06/2026\n'
        'Lưu ý: Báo cáo phục vụ tham khảo. Quyết định đầu tư cần thẩm định độc lập.'
    )
    r_footer.font.size = Pt(9)
    r_footer.italic = True
    r_footer.font.color.rgb = RGBColor.from_string('AAB0B6')

    return doc


# ── Main ──────────────────────────────────────────────────────────────────

def main():
    out_dir = Path(__file__).parent.parent / 'outputs'
    out_dir.mkdir(exist_ok=True)
    out_path = out_dir / 'BaoCao_NganHang_BDS_PhiaNam_THUCTE_20260605.docx'
    print('Dang tao bao cao Word thuc te...')
    doc = build_report()
    doc.save(str(out_path))
    print(f'Xong! {out_path.name}')


if __name__ == '__main__':
    main()
