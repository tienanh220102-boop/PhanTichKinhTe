"""
Tạo báo cáo Word: Thị trường Ngân hàng & BĐS Phía Nam
Chạy: python scripts/tao_bao_cao_word.py
Output: outputs/BaoCao_NganHang_BDS_PhiaNam_YYYYMMDD.docx
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
from pathlib import Path


# ── Helpers ────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_heading(doc, text, level=1, color='1A3A5C'):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = RGBColor.from_string(color)
    return p


def add_para(doc, text, bold=False, size=11, color=None, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return p


def add_bullet(doc, text, size=11):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p


def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'D5E8F5')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


# ── Bảng tổng quan tín hiệu ──────────────────────────────────────────────

def build_signal_table(doc, data):
    """
    data: list of (label, value, color_hex)
    """
    table = doc.add_table(rows=1, cols=len(data))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    row = table.rows[0]
    for i, (label, value, color) in enumerate(data):
        cell = row.cells[i]
        set_cell_bg(cell, color)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run(value + '\n')
        r1.bold = True
        r1.font.size = Pt(13)
        r1.font.color.rgb = RGBColor(255, 255, 255)
        r2 = p.add_run(label)
        r2.font.size = Pt(9)
        r2.font.color.rgb = RGBColor(220, 235, 250)

    # Chiều cao hàng
    for cell in row.cells:
        tc = cell._tc
        trPr = tc.getparent().get_or_add_trPr()
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), '700')
        trPr.append(trHeight)

    return table


# ── Bảng dữ liệu ─────────────────────────────────────────────────────────

def build_data_table(doc, headers, rows_data, col_widths=None):
    table = doc.add_table(rows=1 + len(rows_data), cols=len(headers))
    table.style = 'Table Grid'

    # Header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, '1A3A5C')
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(255, 255, 255)

    # Data rows
    for r_idx, row_data in enumerate(rows_data):
        row = table.rows[r_idx + 1]
        bg = 'EBF5FB' if r_idx % 2 == 0 else 'FFFFFF'
        for c_idx, cell_val in enumerate(row_data):
            cell = row.cells[c_idx]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(cell_val))
            run.font.size = Pt(10)

    # Column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table


# ── Nội dung report ──────────────────────────────────────────────────────

def build_report():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.0)

    # Default font
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal'].font.size = Pt(11)

    now = datetime.now()
    date_str = now.strftime('%d/%m/%Y')

    # ── TIÊU ĐỀ ──────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run('BÁO CÁO THỊ TRƯỜNG')
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor.from_string('1A3A5C')

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(4)
    r2 = p2.add_run('NGÂN HÀNG & BẤT ĐỘNG SẢN PHÍA NAM')
    r2.bold = True
    r2.font.size = Pt(16)
    r2.font.color.rgb = RGBColor.from_string('2471A3')

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_after = Pt(14)
    r3 = p3.add_run(f'Ngày {date_str}  |  TP.HCM · Bình Dương · Đồng Nai · Long An · Bà Rịa-Vũng Tàu')
    r3.font.size = Pt(10)
    r3.font.color.rgb = RGBColor.from_string('7F8C8D')
    r3.italic = True

    add_divider(doc)

    # ── BẢNG TÍN HIỆU TỔNG QUAN ──────────────────────────────────
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    signal_data = [
        ('Xu hướng Lãi suất',  '[TĂNG / GIẢM / ỔN ĐỊNH]',  '2471A3'),
        ('Tín dụng BĐS',       '[NỚI LỎNG / THẮT CHẶT]',   '1A5276'),
        ('Tín hiệu Thị trường','[TÍCH CỰC / TRUNG LẬP / THẬN TRỌNG]', '117A65'),
        ('Số bài phân tích',   '[X bài hôm nay]',           '6C3483'),
    ]
    build_signal_table(doc, signal_data)
    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # ── I. TỔNG QUAN ─────────────────────────────────────────────
    add_heading(doc, 'I. TỔNG QUAN THỊ TRƯỜNG', level=1)

    add_para(doc, 'Tóm tắt điều kiện thị trường ngày hôm nay:', bold=True)
    add_bullet(doc, '[Điểm nổi bật 1: Ví dụ — Lãi suất cho vay mua nhà tại các NHTM lớn tiếp tục ổn định quanh mức X%/năm]')
    add_bullet(doc, '[Điểm nổi bật 2: Ví dụ — Tín dụng BĐS được nới room thêm X nghìn tỷ]')
    add_bullet(doc, '[Điểm nổi bật 3: Ví dụ — Thanh khoản phân khúc căn hộ TP.HCM có dấu hiệu cải thiện]')
    add_bullet(doc, '[Điểm nổi bật 4]')

    doc.add_paragraph()

    # ── II. LÃI SUẤT ─────────────────────────────────────────────
    add_heading(doc, 'II. XU HƯỚNG LÃI SUẤT', level=1)

    add_para(doc, 'A. Lãi suất huy động', bold=True, size=11, color='1A3A5C')

    headers = ['Ngân hàng', 'Kỳ hạn 6 tháng', 'Kỳ hạn 12 tháng', 'Kỳ hạn 24 tháng', 'So tháng trước']
    rows = [
        ['Vietcombank', '[X]%', '[X]%', '[X]%', '[▲/▼/→]'],
        ['BIDV',        '[X]%', '[X]%', '[X]%', '[▲/▼/→]'],
        ['VietinBank',  '[X]%', '[X]%', '[X]%', '[▲/▼/→]'],
        ['Techcombank', '[X]%', '[X]%', '[X]%', '[▲/▼/→]'],
        ['VPBank',      '[X]%', '[X]%', '[X]%', '[▲/▼/→]'],
        ['MB Bank',     '[X]%', '[X]%', '[X]%', '[▲/▼/→]'],
        ['Sacombank',   '[X]%', '[X]%', '[X]%', '[▲/▼/→]'],
        ['ACB',         '[X]%', '[X]%', '[X]%', '[▲/▼/→]'],
    ]
    build_data_table(doc, headers, rows, col_widths=[4.5, 3, 3, 3, 3])

    doc.add_paragraph()
    add_para(doc, 'B. Lãi suất cho vay mua BĐS', bold=True, size=11, color='1A3A5C')

    headers2 = ['Ngân hàng', 'Ưu đãi năm đầu', 'Sau ưu đãi', 'LTV tối đa', 'Thời hạn tối đa']
    rows2 = [
        ['Vietcombank', '[X]%', '[X]%', '[X]%', '[X] năm'],
        ['BIDV',        '[X]%', '[X]%', '[X]%', '[X] năm'],
        ['VietinBank',  '[X]%', '[X]%', '[X]%', '[X] năm'],
        ['Techcombank', '[X]%', '[X]%', '[X]%', '[X] năm'],
        ['VPBank',      '[X]%', '[X]%', '[X]%', '[X] năm'],
    ]
    build_data_table(doc, headers2, rows2, col_widths=[4.5, 3, 3, 3, 3])

    doc.add_paragraph()
    add_para(doc, 'Nhận định xu hướng lãi suất:', bold=True)
    add_bullet(doc, '[Nhận định 1: Ví dụ — Lãi suất huy động có xu hướng tăng nhẹ khi NHNN siết thanh khoản]')
    add_bullet(doc, '[Nhận định 2: Ví dụ — Lãi suất cho vay mua nhà kỳ vọng điều chỉnh trong Q3/2026]')
    add_bullet(doc, '[Nhận định 3]')

    doc.add_paragraph()

    # ── III. TÍN DỤNG BĐS ────────────────────────────────────────
    add_heading(doc, 'III. TÍN DỤNG BẤT ĐỘNG SẢN', level=1)

    add_para(doc, 'A. Chính sách tín dụng BĐS hiện hành', bold=True, size=11, color='1A3A5C')
    add_bullet(doc, '[Chính sách 1: Ví dụ — NHNN nới room tín dụng BĐS thêm X% trong Q2/2026]')
    add_bullet(doc, '[Chính sách 2: Ví dụ — Hệ số rủi ro cho vay BĐS cao cấp điều chỉnh từ X% lên Y%]')
    add_bullet(doc, '[Chính sách 3]')

    doc.add_paragraph()
    add_para(doc, 'B. Tăng trưởng tín dụng BĐS', bold=True, size=11, color='1A3A5C')

    headers3 = ['Chỉ tiêu', 'Cùng kỳ 2025', 'Q1/2026', 'Hiện tại', 'Mục tiêu năm']
    rows3 = [
        ['Tăng trưởng tín dụng BĐS (%)', '[X]%', '[X]%', '[X]%', '[X]%'],
        ['Tỷ trọng tín dụng BĐS / tổng tín dụng', '[X]%', '[X]%', '[X]%', '[X]%'],
        ['Nợ xấu BĐS (%)',                         '[X]%', '[X]%', '[X]%', '[X]%'],
        ['Tín dụng nhà ở xã hội (nghìn tỷ)',       '[X]',  '[X]',  '[X]',  '[X]'],
    ]
    build_data_table(doc, headers3, rows3, col_widths=[5.5, 2.8, 2.8, 2.8, 2.8])

    doc.add_paragraph()

    # ── IV. THỊ TRƯỜNG PHÍA NAM ───────────────────────────────────
    add_heading(doc, 'IV. PHÂN TÍCH THỊ TRƯỜNG PHÍA NAM', level=1)

    khu_vuc = [
        ('TP.HCM', '1A3A5C'),
        ('Bình Dương', '117864'),
        ('Đồng Nai', '6E2F8D'),
        ('Long An', '935116'),
        ('Bà Rịa - Vũng Tàu', '922B21'),
    ]

    for kv, color in khu_vuc:
        add_para(doc, f'  {kv}', bold=True, size=11, color=color)
        add_bullet(doc, f'[Phân khúc nổi bật: Ví dụ — Căn hộ trung cấp {kv} giao dịch X sản phẩm/tháng]')
        add_bullet(doc, f'[Giá trung bình: [X] triệu/m² — tăng/giảm X% so cùng kỳ]')
        add_bullet(doc, '[Điểm đáng chú ý: Ví dụ — Dự án mới mở bán / hạ tầng kết nối]')
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── V. TIN TỨC NỔI BẬT ───────────────────────────────────────
    add_heading(doc, 'V. TIN TỨC NỔI BẬT HÔM NAY', level=1)

    headers4 = ['Tiêu đề tin tức', 'Nguồn', 'Lãi suất', 'Tín dụng BĐS', 'Tác động NĐT phía Nam', 'Khuyến nghị']
    rows4 = [
        ['[Tiêu đề tin 1]', '[Nguồn]', '[TĂNG/GIẢM/ỔN ĐỊNH]', '[NỚI/THẮT]', '[Tác động ngắn gọn]', '[MUA/GIỮ/CHỜ]'],
        ['[Tiêu đề tin 2]', '[Nguồn]', '[TĂNG/GIẢM/ỔN ĐỊNH]', '[NỚI/THẮT]', '[Tác động ngắn gọn]', '[MUA/GIỮ/CHỜ]'],
        ['[Tiêu đề tin 3]', '[Nguồn]', '[TĂNG/GIẢM/ỔN ĐỊNH]', '[NỚI/THẮT]', '[Tác động ngắn gọn]', '[MUA/GIỮ/CHỜ]'],
        ['[Tiêu đề tin 4]', '[Nguồn]', '[TĂNG/GIẢM/ỔN ĐỊNH]', '[NỚI/THẮT]', '[Tác động ngắn gọn]', '[MUA/GIỮ/CHỜ]'],
        ['[Tiêu đề tin 5]', '[Nguồn]', '[TĂNG/GIẢM/ỔN ĐỊNH]', '[NỚI/THẮT]', '[Tác động ngắn gọn]', '[MUA/GIỮ/CHỜ]'],
    ]
    build_data_table(doc, headers4, rows4, col_widths=[5, 2, 2.2, 2.2, 4.5, 2])

    doc.add_paragraph()

    # ── VI. PHÂN TÍCH TỔNG QUAN ───────────────────────────────────
    add_heading(doc, 'VI. PHÂN TÍCH TỔNG QUAN', level=1)

    sections_analysis = [
        ('1. Xu hướng Lãi suất',
         '[Lãi suất huy động/cho vay đang đi theo hướng nào? Tác động đến chi phí vốn của NĐT BĐS? '
         'Dự báo trong 3-6 tháng tới?]'),
        ('2. Tín dụng BĐS',
         '[Chính sách tín dụng BĐS có gì thay đổi? Room tín dụng nới ra hay thắt lại? '
         'Phân khúc nào được ưu tiên (nhà ở xã hội, trung cấp, cao cấp)?]'),
        ('3. Thị trường Phía Nam',
         '[Tác động cụ thể đến TP.HCM, Bình Dương, Đồng Nai, Long An, Bà Rịa-Vũng Tàu? '
         'Phân khúc nào đang hút dòng tiền? Dự án nào đáng quan tâm?]'),
        ('4. Rủi ro cần theo dõi',
         '[Các rủi ro chính: tỷ giá, nợ xấu, pháp lý dự án, thanh khoản thị trường thứ cấp...]'),
    ]

    for title, placeholder in sections_analysis:
        add_para(doc, title, bold=True, size=11, color='1A3A5C')
        add_para(doc, placeholder, size=11, color='7F8C8D')
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    add_divider(doc)

    # ── VII. KHUYẾN NGHỊ ──────────────────────────────────────────
    add_heading(doc, 'VII. KHUYẾN NGHỊ', level=1)

    add_para(doc, 'Dành cho nhà đầu tư BĐS phía Nam hôm nay:', bold=True)
    doc.add_paragraph()

    khuyen_nghi = [
        ('🟢 NÊN LÀM (Cơ hội)', [
            '[Khuyến nghị 1: Ví dụ — Cân nhắc chốt vay mua nhà ngay nếu đang có nhu cầu thực, lãi suất ưu đãi còn hiệu lực]',
            '[Khuyến nghị 2: Ví dụ — Phân khúc đất nền Bình Dương, Long An tiếp tục có cơ sở tăng giá khi hạ tầng hoàn thiện]',
            '[Khuyến nghị 3]',
        ]),
        ('🟡 THEO DÕI (Trung lập)', [
            '[Điểm cần quan sát 1: Ví dụ — Theo dõi chính sách lãi suất điều hành của NHNN trong tháng tới]',
            '[Điểm cần quan sát 2: Ví dụ — Quan sát thanh khoản phân khúc căn hộ cao cấp TP.HCM]',
        ]),
        ('🔴 THẬN TRỌNG (Rủi ro)', [
            '[Rủi ro 1: Ví dụ — Tránh đòn bẩy cao trong bối cảnh lãi suất có thể tăng trở lại]',
            '[Rủi ro 2: Ví dụ — Kiểm tra kỹ pháp lý trước khi xuống tiền tại các dự án mới mở bán]',
        ]),
    ]

    for section_title, items in khuyen_nghi:
        add_para(doc, section_title, bold=True, size=11)
        for item in items:
            add_bullet(doc, item)
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── FOOTER ───────────────────────────────────────────────────
    add_divider(doc)
    p_footer = doc.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_footer.paragraph_format.space_before = Pt(8)
    r_footer = p_footer.add_run(
        f'Báo cáo nội bộ · Ngân Hàng & BĐS Phía Nam · Ngày {date_str}\n'
        'Nguồn tham khảo: VnExpress, CafeF, Thanh Niên, VietnamNet, VOV · Phân tích: AI Agent'
    )
    r_footer.font.size = Pt(9)
    r_footer.italic = True
    r_footer.font.color.rgb = RGBColor.from_string('AAB0B6')

    return doc


# ── Main ──────────────────────────────────────────────────────────────────

def main():
    out_dir = Path(__file__).parent.parent / 'outputs'
    out_dir.mkdir(exist_ok=True)

    date_str = datetime.now().strftime('%Y%m%d')
    out_path = out_dir / f'BaoCao_NganHang_BDS_PhiaNam_{date_str}.docx'

    print('Dang tao bao cao Word...')
    doc = build_report()
    doc.save(str(out_path))
    print(f'Xong! File: {out_path}'.encode('utf-8', errors='replace').decode('utf-8'))


if __name__ == '__main__':
    main()
